import json
import logging
from datetime import datetime
from typing import Literal, Optional

from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel, Field, ValidationError
from sqlalchemy.orm import Session

from app.api.settings import AISettings, get_user_settings
from app.database import get_db
from app.models.document import Document
from app.models.user import User
from app.schemas.document import DocumentContent
from app.services import curriculum_context, images as image_service, providers
from app.services.normalize import collect_image_words, normalize_document
from app.utils.auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/constructor", tags=["constructor"])

SYSTEM_PROMPT = """Eres un experto en educación chilena y pedagogía.
Generas material educativo de alta calidad, adaptado al currículo nacional chileno.
Escribes siempre en español de Chile, con lenguaje claro y apropiado al nivel indicado.
Respondes exclusivamente con un JSON que respete el esquema entregado."""

DocType = Literal["prueba", "evaluacion", "guia", "planificacion", "ficha"]

# Requisito de la licencia CC BY-NC-SA de los pictogramas.
ATRIBUCION_IMAGENES = (
    "Pictogramas: ARASAAC (arasaac.org), Gobierno de Aragón, "
    "autoría de Sergio Palao, licencia CC BY-NC-SA."
)

DOC_TYPE_MAPPING: dict[str, str] = {
    "prueba": "prueba de conocimientos",
    "evaluacion": "evaluación sumativa",
    "guia": "guía de trabajo",
    "planificacion": "planificación de clases",
    "ficha": "ficha de actividades",
}


# ---------------------------------------------------------------------------
# Modelos de entrada
# ---------------------------------------------------------------------------

class GenerateRequest(BaseModel):
    doc_type: DocType
    subject: str = Field(..., min_length=1, max_length=120)
    grade_level: str = Field(..., min_length=1, max_length=60)
    topic: str = Field(..., min_length=1, max_length=300)
    instructions: str = Field("", max_length=4000)
    num_questions: int = Field(10, ge=1, le=30)
    difficulty: Literal["fácil", "medio", "difícil"] = "medio"
    include_images: bool = False
    include_answers: bool = True
    oa_codes: list[str] = Field(default_factory=list, max_length=20)
    indicator_refs: list[str] = Field(default_factory=list, max_length=300)
    provider: Optional[str] = None  # openrouter | gemini | openai | xai | auto


class DocumentSave(BaseModel):
    title: str = Field(..., min_length=1, max_length=300)
    doc_type: DocType
    subject: str = ""
    grade_level: str = ""
    content: dict = Field(default_factory=dict)
    ai_prompt: str = ""
    images: list = Field(default_factory=list)
    activity_images: dict = Field(default_factory=dict)


class SaveToCalendarRequest(DocumentSave):
    title: str = ""
    event_date: str


class OptimizeRequest(BaseModel):
    instructions: str = Field(..., min_length=1, max_length=4000)
    doc_type: str = ""
    subject: str = ""
    grade_level: str = ""
    topic: str = ""
    oa_codes: list[str] = Field(default_factory=list, max_length=20)
    indicator_refs: list[str] = Field(default_factory=list, max_length=300)


class SearchImagesRequest(BaseModel):
    words: list[str] = Field(default_factory=list, max_length=40)
    style: Literal["photo", "coloring"] = "photo"


class ExportRequest(BaseModel):
    content: dict
    activity_images: dict[str, str] = Field(default_factory=dict)


# ---------------------------------------------------------------------------
# Construcción del prompt
# ---------------------------------------------------------------------------

def build_prompt(req: GenerateRequest, curriculum_block: str = "") -> str:
    doc_name = DOC_TYPE_MAPPING.get(req.doc_type, req.doc_type)
    parts = [
        f"Crea una {doc_name} completa para la asignatura de {req.subject},",
        f"nivel {req.grade_level}, sobre el tema: '{req.topic}'.",
        f"Dificultad: {req.difficulty}.",
    ]

    if req.doc_type in ("prueba", "evaluacion"):
        parts.append(
            f"Incluye {req.num_questions} preguntas variadas "
            "(selección múltiple, verdadero/falso y desarrollo)."
        )
        if req.include_answers:
            parts.append(
                "Agrega una sección final de tipo 'answers' con la pauta de corrección: "
                "cada ítem con su respuesta correcta y su puntaje."
            )
    elif req.doc_type == "guia":
        parts.append(
            f"Estructura el documento en bloques de ACTIVIDADES (sections type='activities', ítems "
            f"type='activity'), con {req.num_questions} actividades en total, graduadas de menor a "
            "mayor dificultad. Esto NO es una prueba: no pongas alternativas ni puntajes salvo que las "
            "instrucciones del docente lo pidan explícitamente."
        )
        parts.append(
            "Cada actividad (ContentItem) llena tres campos propios además de 'text': 'purpose' con el "
            "propósito pedagógico en una frase breve (qué habilidad ejercita y por qué); 'text' con la "
            "instrucción para el estudiante, una frase simple y directa, apropiada a 1°-2° básico; e "
            "'indicator_ref' con la referencia del indicador de evaluación que esa actividad trabaja, "
            "copiada EXACTA (sin corchetes) desde el bloque de currículum de abajo — nunca inventada, y "
            "vacía si el OA correspondiente no trae indicadores."
        )
        parts.append(
            "ACTIVIDADES CON IMÁGENES: cuando una actividad requiera que el alumno observe dibujos, "
            "pon las palabras en el campo 'image_words' del ítem. Deben ser sustantivos concretos y "
            "fáciles de ilustrar (animales, objetos, frutas). Nunca uses una sola letra como palabra. "
            "El texto del ítem NO debe listar esas palabras ni describir las imágenes: nada de "
            "'Imagen de una abeja' ni viñetas con los dibujos. Si la actividad pide pintar o colorear, "
            "marca image_style='coloring'. Para 1° y 2° básico, casi toda actividad debería traer al "
            "menos una palabra ilustrable: es el apoyo visual que distingue a la guía de una prueba."
        )
    elif req.doc_type == "planificacion":
        parts.append(
            "Incluye secciones para: objetivos de aprendizaje, habilidades, actitudes, "
            "momentos de inicio-desarrollo-cierre, y evaluación."
        )
    elif req.doc_type == "ficha":
        parts.append("Incluye datos del estudiante, una sección de contenidos y actividades prácticas.")

    if req.instructions.strip():
        parts.append(f"Instrucciones adicionales del docente: {req.instructions.strip()}")

    parts.append(
        "FORMATO: cada sección usa 'body' para texto corrido o 'items' para listas de preguntas y "
        "actividades, nunca ambos para el mismo contenido. El campo 'text' de cada ítem no lleva "
        "el número al inicio: el número va en el campo 'number'."
    )

    prompt = " ".join(parts)
    # El contexto curricular va en el mensaje del usuario, no en el system:
    # en el system el modelo lo trata como rol a interpretar, no como dato a citar.
    return prompt + curriculum_block


# ---------------------------------------------------------------------------
# Generación con validación y reintento
# ---------------------------------------------------------------------------

async def _generate_validated(
    settings: AISettings,
    prompt: str,
    provider: Optional[str],
) -> tuple[DocumentContent, providers.GenerationResult]:
    """Genera y valida. Un reintento con el error como feedback, y basta.

    Antes no había validación alguna: `response_mime_type="application/json"`
    garantiza JSON sintácticamente válido, no que tenga la forma pedida. Cuando
    el modelo devolvía secciones en texto plano, el pipeline de imágenes quedaba
    inerte y nadie se enteraba.
    """
    result = await providers.generate_json(settings, prompt=prompt, system=SYSTEM_PROMPT, provider=provider)

    try:
        return normalize_document(result.content), result
    except (ValidationError, ValueError) as exc:
        # `except ... as exc` borra la variable al salir del bloque, así que el
        # mensaje se copia antes de usarlo en el prompt de reintento.
        first_error = str(exc)
        logger.warning("Documento inválido en el primer intento: %s", first_error)

    retry_prompt = (
        f"{prompt}\n\n"
        "El intento anterior no respetó el esquema. Error de validación:\n"
        f"{first_error}\n"
        "Corrige la estructura y responde de nuevo solo con el JSON válido."
    )
    retry = await providers.generate_json(
        settings, prompt=retry_prompt, system=SYSTEM_PROMPT, provider=provider
    )
    retry.cost += result.cost

    try:
        return normalize_document(retry.content), retry
    except (ValidationError, ValueError) as second_error:
        # Falla visible: antes esto degradaba en silencio y la profesora recibía
        # un documento inservible sin ninguna señal de que algo salió mal.
        raise HTTPException(
            status_code=502,
            detail=(
                "El modelo devolvió un documento con estructura inválida dos veces seguidas. "
                "Prueba de nuevo o cambia de proveedor en Configuración."
            ),
        ) from second_error


def _resolve_oa_codes(doc: DocumentContent, requested: list[str], allowed: set[str]) -> list[str]:
    """Fija los OA del documento y devuelve los que el modelo haya inventado.

    Si la profesora eligió OA explícitamente, esos mandan: son la fuente de
    verdad y no depende de que el modelo se acuerde de repetirlos en el JSON
    (en la práctica, muchas veces devuelve la lista vacía). Sin selección
    explícita se conserva lo que propuso el modelo, descartando lo inventado.
    """
    if not allowed:
        doc.metadata.oa_codes = []
        return []

    if requested:
        doc.metadata.oa_codes = [c.strip().upper() for c in requested if c.strip().upper() in allowed]
        return []

    kept, invented = [], []
    for code in doc.metadata.oa_codes:
        (kept if code.strip().upper() in allowed else invented).append(code)
    doc.metadata.oa_codes = kept
    if invented:
        logger.warning("El modelo inventó códigos OA inexistentes y se descartaron: %s", invented)
    return invented


def _resolve_indicator_refs(doc: DocumentContent, allowed: set[str]) -> list[str]:
    """Descarta indicator_ref inventados: solo puede citar lo que el contexto trajo.

    Igual criterio que `_resolve_oa_codes`: el indicador que ve la profesora en
    el documento tiene que ser uno oficial de verdad, no algo que el modelo se
    inventó al copiar el formato.
    """
    invented = []
    for section in doc.sections:
        for item in section.items:
            if not item.indicator_ref:
                continue
            if item.indicator_ref.strip().upper() not in allowed:
                invented.append(item.indicator_ref)
                item.indicator_ref = ""
    if invented:
        logger.warning("El modelo devolvió indicator_ref inexistentes y se descartaron: %s", invented)
    return invented


@router.post("/generate")
async def generate_document(
    req: GenerateRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    settings = get_user_settings(current_user.id, db)

    ctx = curriculum_context.build_context(
        db, req.grade_level, req.subject, req.oa_codes, req.indicator_refs
    )
    prompt = build_prompt(req, ctx.block)

    doc, result = await _generate_validated(settings, prompt, req.provider)

    allowed = curriculum_context.valid_codes(db, req.grade_level, req.subject)
    _resolve_oa_codes(doc, req.oa_codes, allowed)
    _resolve_indicator_refs(doc, ctx.indicator_refs)

    if not doc.metadata.subject:
        doc.metadata.subject = req.subject
    if not doc.metadata.grade:
        doc.metadata.grade = req.grade_level
    if not doc.metadata.topic:
        doc.metadata.topic = req.topic

    logger.info(
        "Documento generado: %s/%s — %d secciones, %d palabras ilustrables, costo $%.5f",
        result.provider,
        result.model,
        len(doc.sections),
        len(collect_image_words(doc)),
        result.cost,
    )

    return {
        "content": doc.model_dump(),
        "prompt": prompt,
        "images": [],
        "provider_used": result.provider,
        "model_used": result.model,
        "cost": result.cost,
        "curriculum_grounded": bool(ctx.block),
    }


# ---------------------------------------------------------------------------
# Imágenes
# ---------------------------------------------------------------------------

class CoverRequest(BaseModel):
    subject: str = ""
    grade_level: str = ""
    topic: str = Field(..., min_length=1, max_length=300)


@router.post("/generate-image")
async def generate_cover_image(
    data: CoverRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    settings = get_user_settings(current_user.id, db)
    if not settings.openrouter_key:
        raise HTTPException(
            status_code=503,
            detail="Configura tu API Key de OpenRouter para generar imágenes.",
        )
    url = await image_service.generate_cover(settings, data.subject, data.grade_level, data.topic)
    if not url:
        raise HTTPException(status_code=502, detail="No se pudo generar la imagen.")
    return {"image_url": url}


@router.post("/search-images")
async def search_images(
    data: SearchImagesRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Genera las ilustraciones de un lote de palabras. Devuelve {palabra: url}."""
    settings = get_user_settings(current_user.id, db)
    words = {w.strip().lower(): data.style for w in data.words if len(w.strip()) > 1}
    if not words:
        return {"images": {}}
    return {"images": await image_service.generate_images(settings, words)}


# ---------------------------------------------------------------------------
# Optimización de instrucciones
# ---------------------------------------------------------------------------

@router.post("/optimize-instructions")
async def optimize_instructions(
    data: OptimizeRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Convierte instrucciones simples en un prompt técnico.

    Una llamada al proveedor por click. La versión anterior llamaba al modelo,
    descartaba el resultado con un `raise ValueError("use_plain")` y volvía a
    llamar: cada click costaba el doble.
    """
    settings = get_user_settings(current_user.id, db)

    # El anclaje curricular solo se inyecta cuando la profesora eligió OA.
    # Sin selección, `build_context` devuelve hasta 12 objetivos y el prompt de
    # salida son 2-4 oraciones: no puede cubrirlos todos sin volverse genérico,
    # que es justo lo contrario de lo que se busca al optimizar.
    curriculum_block = ""
    if data.oa_codes and data.grade_level and data.subject:
        curriculum_block = curriculum_context.build_context(
            db, data.grade_level, data.subject, data.oa_codes, data.indicator_refs
        ).block

    optimize_prompt = f"""Eres un experto en ingeniería de prompts y diseño instruccional.

CONTEXTO DEL DOCUMENTO A GENERAR:
- Tipo: {data.doc_type}
- Asignatura: {data.subject}
- Nivel escolar: {data.grade_level}
- Tema: {data.topic}

INSTRUCCIONES SIMPLES DEL DOCENTE:
\"\"\"{data.instructions.strip()}\"\"\"

TU TAREA:
Transforma esas instrucciones en un prompt técnico y específico que extraiga el máximo
del modelo al generar el documento. No parafrasees: expande y enriquece con
- nivel cognitivo explícito según la taxonomía de Bloom;
- especificaciones pedagógicas concretas (tipo de preguntas, enfoque didáctico, habilidades);
- restricciones de formato que mejoren la calidad (longitud, tono, complejidad léxica del nivel);
- criterios de calidad propios de ese tipo de documento;
- consideraciones disciplinares relevantes para {data.subject}.

FORMATO DE RESPUESTA: un bloque continuo de 2 a 4 oraciones, en español, sin títulos,
sin viñetas y sin comillas. Directo y técnico, listo para usarse como instrucción adicional."""

    if curriculum_block:
        optimize_prompt += curriculum_block
        optimize_prompt += (
            "\nUSO DE LOS OBJETIVOS DE APRENDIZAJE:\n"
            "Las instrucciones que escribas deben apuntar específicamente a los OA listados: "
            "usa sus verbos y su nivel de exigencia para fijar qué habilidad se evalúa y con "
            "qué profundidad. NO cites los códigos OA en tu respuesta ni menciones el "
            "currículum: el texto se inserta en un campo de instrucciones adicionales, no en "
            "el documento que verá el alumno.\n"
        )

    result = await providers.generate_text(settings, prompt=optimize_prompt, max_tokens=2000)
    optimized = result.text.strip()
    if not optimized:
        raise HTTPException(status_code=502, detail="El modelo no devolvió instrucciones optimizadas.")
    return {"optimized": optimized, "provider_used": result.provider, "cost": result.cost}


# ---------------------------------------------------------------------------
# Persistencia
# ---------------------------------------------------------------------------

def _build_document(user_id: int, data: DocumentSave, title: str) -> Document:
    return Document(
        user_id=user_id,
        title=title,
        doc_type=data.doc_type,
        subject=data.subject or None,
        grade_level=data.grade_level or None,
        content=data.content,
        ai_prompt=data.ai_prompt or None,
        images=data.images or [],
    )


@router.post("/save", response_model=dict)
async def save_document(
    doc: DocumentSave,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    db_doc = _build_document(current_user.id, doc, doc.title)
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    return {"id": db_doc.id, "message": "Documento guardado"}


CATEGORY_MAP = {
    "prueba": "evaluacion",
    "evaluacion": "evaluacion",
    "planificacion": "planificacion",
    "guia": "general",
    "ficha": "general",
}
COLOR_MAP = {"evaluacion": "#EF4444", "planificacion": "#10B981", "general": "#6B7280"}


@router.post("/save-to-calendar", response_model=dict)
async def save_to_calendar(
    data: SaveToCalendarRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from app.models.event import Event

    title = data.title or f"{data.doc_type} - {data.subject}".strip(" -")
    db_doc = _build_document(current_user.id, data, title)
    db.add(db_doc)
    db.flush()

    category = CATEGORY_MAP.get(data.doc_type, "general")
    try:
        event_dt = datetime.fromisoformat(data.event_date)
    except ValueError:
        event_dt = datetime.utcnow()

    db_event = Event(
        user_id=current_user.id,
        title=db_doc.title,
        description=f"[doc_id:{db_doc.id}] {data.subject} - {data.grade_level}".strip(" -"),
        start_datetime=event_dt,
        end_datetime=event_dt,
        all_day=True,
        color=COLOR_MAP.get(category, "#6B7280"),
        category=category,
    )
    db.add(db_event)
    db.commit()
    db.refresh(db_doc)
    db.refresh(db_event)

    return {
        "doc_id": db_doc.id,
        "event_id": db_event.id,
        "message": "Documento guardado y agregado al calendario",
    }


@router.get("/documents")
def get_documents(
    doc_type: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    query = db.query(Document).filter(Document.user_id == current_user.id)
    if doc_type:
        query = query.filter(Document.doc_type == doc_type)
    return [
        {
            "id": d.id,
            "title": d.title,
            "doc_type": d.doc_type,
            "subject": d.subject,
            "grade_level": d.grade_level,
            "created_at": d.created_at.isoformat(),
        }
        for d in query.order_by(Document.created_at.desc()).all()
    ]


@router.get("/documents/{doc_id}")
def get_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    # Los documentos guardados antes del esquema validado tienen `content` con
    # secciones en texto plano; se normalizan al vuelo para que abran igual.
    content = doc.content
    try:
        content = normalize_document(content).model_dump() if content else None
    except Exception:
        logger.warning("No se pudo normalizar el documento %s; se devuelve tal cual", doc_id)

    return {
        "id": doc.id,
        "title": doc.title,
        "doc_type": doc.doc_type,
        "subject": doc.subject,
        "grade_level": doc.grade_level,
        "content": content,
        "raw_html": doc.raw_html,
        "images": doc.images,
        "ai_prompt": doc.ai_prompt,
        "created_at": doc.created_at.isoformat(),
    }


@router.delete("/documents/{doc_id}")
def delete_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    db.delete(doc)
    db.commit()
    return {"message": "Documento eliminado"}


# ---------------------------------------------------------------------------
# Exportación
# ---------------------------------------------------------------------------

def _strip_md(text) -> str:
    import re

    if not text:
        return ""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", str(text))
    return re.sub(r"\*(.+?)\*", r"\1", text).strip()


@router.post("/export-pdf")
async def export_pdf(data: ExportRequest, current_user: User = Depends(get_current_user)):
    import io
    import os

    from fastapi.responses import StreamingResponse
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        HRFlowable,
        Image as RLImage,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    doc_content = normalize_document(data.content)
    title = _strip_md(doc_content.title) or "Documento"

    def image_flowable(word: str, size_mm: float = 24):
        url = data.activity_images.get(word.lower().strip(), "")
        if not url.startswith("/static/"):
            return None
        path = "/app" + url
        if not os.path.exists(path):
            return None
        try:
            return RLImage(path, width=size_mm * mm, height=size_mm * mm)
        except Exception:
            return None

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=18 * mm, rightMargin=18 * mm, topMargin=15 * mm, bottomMargin=15 * mm,
    )
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(title.upper(), ParagraphStyle(
        "DocTitle", parent=styles["Title"], fontSize=16, alignment=TA_CENTER, spaceAfter=4)))

    meta = doc_content.metadata
    meta_parts = []
    if meta.subject:
        meta_parts.append(f"<b>Asignatura:</b> {meta.subject}")
    if meta.grade:
        meta_parts.append(f"<b>Nivel:</b> {meta.grade}")
    if meta.total_points:
        meta_parts.append(f"<b>Puntaje:</b> {meta.total_points} pts")
    if meta.oa_codes:
        meta_parts.append(f"<b>OA:</b> {', '.join(meta.oa_codes)}")
    if meta_parts:
        story.append(Paragraph("   |   ".join(meta_parts), ParagraphStyle(
            "Meta", parent=styles["Normal"], fontSize=10, alignment=TA_CENTER, spaceAfter=6)))

    story.append(HRFlowable(width="100%", thickness=1, color=colors.black, spaceAfter=6))
    header = Table(
        [["Nombre: ________________________", "Curso: ______________", "Fecha: ______________"]],
        colWidths=[90 * mm, 50 * mm, 50 * mm],
    )
    header.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(header)
    story.append(Spacer(1, 6 * mm))

    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, spaceAfter=4, leading=14)

    if doc_content.instructions:
        story.append(Paragraph(
            f"<b>Instrucciones:</b> {_strip_md(doc_content.instructions)}",
            ParagraphStyle("Inst", parent=body_style, backColor=colors.HexColor("#f9f9f9"),
                           borderPadding=6, spaceAfter=8)))

    for section in doc_content.sections:
        if section.title:
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey, spaceAfter=2))
            story.append(Paragraph(_strip_md(section.title), ParagraphStyle(
                "SecH", parent=styles["Heading2"], fontSize=11, spaceAfter=4)))
        if section.body:
            story.append(Paragraph(_strip_md(section.body), body_style))

        for item in section.items:
            line = f"{item.number}. {_strip_md(item.text)}"
            if item.points:
                line += f"  <font size='8' color='grey'>({item.points} pts)</font>"
            story.append(Paragraph(line, ParagraphStyle(
                "Q", parent=body_style, backColor=colors.HexColor("#f5f5f5"), borderPadding=5)))

            if item.image_words:
                cells, labels = [], []
                for word in item.image_words:
                    flowable = image_flowable(word)
                    cells.append(flowable if flowable else Paragraph("", body_style))
                    labels.append(Paragraph(
                        f"<para align='center'><font size='8'>{word.capitalize()}</font></para>", body_style))
                table = Table([cells, labels], colWidths=[28 * mm] * len(cells))
                table.setStyle(TableStyle([
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]))
                story.append(table)
                story.append(Spacer(1, 3 * mm))

            for index, option in enumerate(item.options):
                story.append(Paragraph(
                    f"&nbsp;&nbsp;&nbsp;{chr(65 + index)}) {_strip_md(option)}", body_style))
            if item.answer:
                story.append(Paragraph(
                    f"<font color='green'>✓ {_strip_md(item.answer)}</font>",
                    ParagraphStyle("Ans", parent=body_style, fontSize=9)))

        story.append(Spacer(1, 4 * mm))

    if data.activity_images:
        # La licencia CC BY-NC-SA de ARASAAC exige citar la fuente en el
        # material donde aparecen los pictogramas.
        story.append(Spacer(1, 4 * mm))
        story.append(Paragraph(
            f"<font size='7' color='grey'>{ATRIBUCION_IMAGENES}</font>",
            ParagraphStyle("Attr", parent=body_style, fontSize=7),
        ))

    pdf.build(story)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{title.replace(" ", "_")[:50]}.pdf"'},
    )


@router.post("/export-docx")
async def export_docx(data: ExportRequest, current_user: User = Depends(get_current_user)):
    import io
    import os

    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    from fastapi.responses import StreamingResponse

    doc_content = normalize_document(data.content)
    title = _strip_md(doc_content.title) or "Documento"

    docx = DocxDocument()
    for section in docx.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    heading = docx.add_heading(title.upper(), level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc_content.metadata
    meta_parts = []
    if meta.subject:
        meta_parts.append(f"Asignatura: {meta.subject}")
    if meta.grade:
        meta_parts.append(f"Nivel: {meta.grade}")
    if meta.total_points:
        meta_parts.append(f"Puntaje: {meta.total_points} pts")
    if meta.oa_codes:
        meta_parts.append(f"OA: {', '.join(meta.oa_codes)}")
    if meta_parts:
        paragraph = docx.add_paragraph("   |   ".join(meta_parts))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(10)

    docx.add_paragraph("─" * 80)
    table = docx.add_table(rows=1, cols=3)
    cells = table.rows[0].cells
    cells[0].text = "Nombre: _______________________"
    cells[1].text = "Curso: _______________"
    cells[2].text = "Fecha: _______________"
    docx.add_paragraph("")

    if doc_content.instructions:
        paragraph = docx.add_paragraph()
        label = paragraph.add_run("Instrucciones: ")
        label.bold = True
        label.font.size = Pt(10)
        body = paragraph.add_run(_strip_md(doc_content.instructions))
        body.font.size = Pt(10)

    for section in doc_content.sections:
        if section.title:
            docx.add_heading(_strip_md(section.title), level=2)
        if section.body:
            docx.add_paragraph(_strip_md(section.body))

        for item in section.items:
            label = f"({item.points} pts)" if item.points else ""
            paragraph = docx.add_paragraph()
            run = paragraph.add_run(f"{item.number}. {_strip_md(item.text)} {label}".rstrip())
            run.font.size = Pt(10)

            if item.image_words:
                img_table = docx.add_table(rows=2, cols=len(item.image_words))
                for column, word in enumerate(item.image_words):
                    url = data.activity_images.get(word.lower().strip(), "")
                    path = "/app" + url if url.startswith("/static/") else None
                    cell = img_table.cell(0, column)
                    if path and os.path.exists(path):
                        try:
                            picture_paragraph = cell.paragraphs[0]
                            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            picture_paragraph.add_run().add_picture(path, width=Inches(0.9))
                        except Exception:
                            cell.text = f"[{word}]"
                    else:
                        cell.text = f"[{word}]"
                    caption = img_table.cell(1, column).paragraphs[0]
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.add_run(word.capitalize()).font.size = Pt(8)
                docx.add_paragraph("")

            for index, option in enumerate(item.options):
                option_paragraph = docx.add_paragraph(f"    {chr(65 + index)}) {_strip_md(option)}")
                option_paragraph.paragraph_format.left_indent = Inches(0.3)
            if item.answer:
                answer_paragraph = docx.add_paragraph(f"✓ {_strip_md(item.answer)}")
                answer_paragraph.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)
                answer_paragraph.runs[0].font.size = Pt(9)

    if data.activity_images:
        attr = docx.add_paragraph(ATRIBUCION_IMAGENES)
        attr.runs[0].font.size = Pt(7)
        attr.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{title.replace(" ", "_")[:50]}.docx"'},
    )
